"""Document analysis — pypdf for text PDFs, Gemini vision fallback for scanned. Open access."""

from __future__ import annotations

import asyncio
import base64
import io
from pathlib import Path

import requests
from astrbot.api.event import AstrMessageEvent, filter
from astrbot.api.message_components import File
from astrbot.api.star import Context, Star
try:
    from xiaoning_runtime import chat_response_content, defer_stop_event
except ImportError:
    from data.plugins.xiaoning_runtime import chat_response_content, defer_stop_event

PROXY = "http://127.0.0.1:3000/v1/chat/completions"
MAX_CHARS = 40_000
MAX_PDF_PAGES = 50
MIN_TEXT_LENGTH = 50  # below this, treat PDF as scanned
SCANNED_RENDER_PAGES = 3

SYSTEM_PROMPT = "你是一个专业的文档分析助手。用中文直接输出分析结果，不要尝试执行代码或调用工具。"


class PdfAnalysis(Star):
    def __init__(self, context: Context, config: dict | None = None):
        super().__init__(context, config)
        self.config = config or {}
    @filter.platform_adapter_type(filter.PlatformAdapterType.ALL, priority=960)
    @defer_stop_event
    async def on_message(self, event: AstrMessageEvent):
        sender_id = str(getattr(event, "get_sender_id", lambda: "")() or "")
        if not sender_id:
            return

        components = getattr(getattr(event, "message_obj", None), "message", None) or []
        files = [
            c for c in (components if isinstance(components, list) else [])
            if isinstance(c, File) or (isinstance(c, dict) and c.get("type") == "file")
        ]

        text = self._msg(event)
        if text.startswith("/analysis") or text.startswith("/分析"):
            prompt_override = text.split(maxsplit=1)[1] if " " in text else ""
        else:
            prompt_override = ""
        has_prompt = bool(prompt_override)

        if not files and not has_prompt:
            return
        if not (event.is_private_chat() or event.is_at_or_wake_command):
            return
        event.stop_event()

        if not files:
            return

        content = ""
        source_name = ""
        user_message: str | list[dict] = ""

        if files:
            file_obj = files[0]
            if isinstance(file_obj, File):
                path_str = (await file_obj.get_file()) or ""
                source_name = file_obj.name or ""
            elif isinstance(file_obj, dict):
                data = file_obj.get("data", {})
                path_str = data.get("file", "") or data.get("path", "") or data.get("url", "")
                source_name = data.get("name", "") or Path(str(path_str)).name
            else:
                path_str = ""
                source_name = ""

            file_path = Path(str(path_str)) if path_str else Path()
            source_name = source_name or file_path.name
            suffix = file_path.suffix.lower()

            if not file_path.is_file():
                yield event.plain_result("文件未找到，请重新发送。")
                return

            # ── PDF: pypdf text first, vision fallback for scanned ──
            if suffix == ".pdf":
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(file_path))
                    pages_text = []
                    total = 0
                    for page in reader.pages[:MAX_PDF_PAGES]:
                        pt = (page.extract_text() or "")[:3000]
                        if pt.strip():
                            pages_text.append(pt)
                            total += len(pt)
                            if total > MAX_CHARS:
                                break
                    content = "\n\n".join(pages_text)
                    page_count = len(reader.pages)
                except Exception as exc:
                    yield event.plain_result(f"PDF 读取失败：{type(exc).__name__}")
                    return

                if len(content.strip()) >= MIN_TEXT_LENGTH:
                    # Text-based PDF — send extracted text to Gemini
                    source_name = f"{file_path.name} ({len(pages_text)} 页文字)"
                    analysis_prompt = prompt_override or (
                        "请分析以下文档内容，用中文直接输出：\n"
                        "1. 主题和核心观点\n"
                        "2. 关键论据或数据\n"
                        "3. 局限性或未解决的问题（如有）\n\n"
                        f"文档内容：\n{content[:MAX_CHARS]}"
                    )
                    user_message = analysis_prompt if not prompt_override else (
                        f"文档内容：\n{content[:MAX_CHARS]}\n\n用户要求：{prompt_override}"
                    )
                else:
                    # Scanned PDF — render first few pages as small PNGs
                    try:
                        import pypdfium2 as pdfium
                        from PIL import Image
                        pdf = pdfium.PdfDocument(str(file_path))
                        render_count = min(len(pdf), SCANNED_RENDER_PAGES)
                        page_pngs = []
                        for i in range(render_count):
                            page = pdf[i]
                            bitmap = page.render(scale=0.6)
                            img = bitmap.to_pil().convert("RGB")
                            img.thumbnail((600, 800), Image.LANCZOS)
                            buf = io.BytesIO()
                            img.save(buf, format="JPEG", quality=60)
                            page_pngs.append(buf.getvalue())
                            page.close()
                            bitmap.close()
                        pdf.close()

                        user_content: list[dict] = [{
                            "type": "text",
                            "text": prompt_override or (
                                f"这份扫描PDF共{len(pdf)}页（显示前{render_count}页），请用中文分析"
                                "其内容、数据和关键信息。"
                            ),
                        }]
                        for img_bytes in page_pngs:
                            img_b64 = base64.b64encode(img_bytes).decode()
                            user_content.append({
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
                            })
                        user_message = user_content
                        source_name = f"{file_path.name} (扫描件, {render_count}页渲染)"
                    except Exception as exc:
                        yield event.plain_result(
                            f"PDF 似乎为扫描件且渲染失败：{type(exc).__name__}。"
                            f"请尝试发送可读取文本的 PDF。"
                        )
                        return

            # ── Images: direct to Gemini vision ──
            elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
                try:
                    img_bytes = file_path.read_bytes()
                    img_b64 = base64.b64encode(img_bytes).decode()
                    mime = "image/png" if suffix == ".png" else "image/jpeg"
                    if suffix == ".webp": mime = "image/webp"
                    if suffix == ".gif": mime = "image/gif"
                    user_message = [
                        {"type": "text", "text": prompt_override or "请描述和分析这张图片的内容，用中文输出。"},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_b64}"}},
                    ]
                except Exception as exc:
                    yield event.plain_result(f"图片读取失败：{type(exc).__name__}")
                    return

            # ── DOCX: python-docx → text ──
            elif suffix == ".docx":
                try:
                    from docx import Document
                    doc = Document(str(file_path))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    content = "\n".join(paragraphs)[:MAX_CHARS]
                except Exception as exc:
                    yield event.plain_result(f"Word 文档读取失败：{type(exc).__name__}")
                    return

            # ── Text files ──
            elif suffix in {".txt", ".md", ".py", ".json", ".csv"}:
                try:
                    content = file_path.read_text(encoding="utf-8", errors="replace")[:MAX_CHARS]
                except Exception:
                    yield event.plain_result("文件读取失败。")
                    return

            else:
                yield event.plain_result(
                    "请发送 PDF / DOCX / TXT / MD / 图片文件。"
                )
                return

        # ── build final prompt for text-based content ──
        if isinstance(user_message, str):
            if not user_message:
                if prompt_override:
                    user_message = f"文档内容：\n{content[:MAX_CHARS]}\n\n用户要求：{prompt_override}" if content else prompt_override
                elif content:
                    user_message = (
                        "请分析以下文档内容，用中文直接输出：\n"
                        "1. 主题和核心观点\n"
                        "2. 关键论据或数据\n"
                        "3. 局限性或未解决的问题（如有）\n\n"
                        f"文档内容：\n{content[:MAX_CHARS]}"
                    )
                else:
                    yield event.plain_result("请发送一个文件或输入分析要求。")
                    return

        yield event.plain_result(
            f"📄 分析中{f'（{source_name}）' if source_name else ''}…"
        )

        try:
            resp = await asyncio.to_thread(
                requests.post,
                PROXY,
                json={
                    "model": "gemini-3.6-flash",
                    "messages": [
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": user_message},
                    ],
                    "max_tokens": 2000,
                },
                timeout=120,
            )
            result = chat_response_content(resp)
        except Exception as exc:
            yield event.plain_result(f"分析服务暂时不可用：{str(exc) or type(exc).__name__}")
            return

        yield event.plain_result(result)

    @staticmethod
    def _msg(event: AstrMessageEvent) -> str:
        getter = getattr(event, "get_message_str", None)
        if callable(getter):
            return str(getter() or "")
        return str(getattr(event, "message_str", "") or "")
